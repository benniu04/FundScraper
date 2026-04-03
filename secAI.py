from sec_api import FullTextSearchApi
from datetime import datetime
import os
import requests
import openai
from config import SEC_IO_API_KEY, OPENAI_API_KEY
from docx import Document

SEC_API_KEY = SEC_IO_API_KEY
ACTUAL_OPENAI_API_KEY = OPENAI_API_KEY

DOC_PATH = "sectest040326.docx"

if os.path.exists(DOC_PATH):
    print(f"Loading existing document: {DOC_PATH}")
    doc = Document(DOC_PATH)
else:
    print(f"Creating new document: {DOC_PATH}")
    doc = Document()
    doc.add_heading('COMPANIES RANKINGS', level=1)

# Set up OpenAI client
if not ACTUAL_OPENAI_API_KEY:
    print("Missing OpenAI API key—analysis won’t work.")
    openai_client = None
else:
    openai_client = openai.OpenAI(api_key=ACTUAL_OPENAI_API_KEY)

# Fixed filing parameters
DEFAULT_FORMS = ["10-K", "10-Q", "13F-HR", "8-K", "D", "S-1"]
DEFAULT_START = datetime(2010, 1, 10)
DEFAULT_END = datetime(2025, 12, 10)
LIMIT = 200


def fetch_filings(query, limit=LIMIT):
    """Fetch the first few SEC filings based on query keywords using sec-api.io."""
    if not SEC_API_KEY or SEC_API_KEY == "YOUR_SEC_API_KEY_HERE":
        print("Please provide a valid sec-api.io API key.")
        return []

    full_text_search_api = FullTextSearchApi(api_key=SEC_API_KEY)
    search_params = {
        "query": query,
        "formTypes": DEFAULT_FORMS,
        "startDate": DEFAULT_START.strftime("%Y-%m-%d"),
        "endDate": DEFAULT_END.strftime("%Y-%m-%d"),
        "page": 1
    }

    try:
        print(f"Fetching filings for query: {query}...")
        response = full_text_search_api.get_filings(search_params)
        total_filings = response["total"]["value"]
        filings = response["filings"][:limit]

        if not filings:
            print(f"No filings found for query: {query}")
            return []

        print(f"Found {len(filings)} filings (out of {total_filings} total matches).")
        return [
            {
                "accession_number": filing["accessionNo"],
                "form_type": filing["formType"],
                "entity_name": filing["companyNameLong"],
                "filing_date": filing["filedAt"].split("T")[0],
                "document_url": filing["filingUrl"]
            }
            for filing in filings
        ]
    except Exception as e:
        print(f"Error fetching filings: {e}")
        return []


def analyze_investment_potential_with_openai(filings, client_name, client_specialization):
    """Use OpenAI to identify companies likely to invest in the client based on filings."""
    if not openai_client:
        return "No OpenAI client available for analysis."

    # Combine content from all filings
    headers = {"User-Agent": "BenNiu (nub38bn@gmail.com)"}
    combined_content = ""
    for filing in filings:
        try:
            response = requests.get(filing["document_url"], headers=headers)
            response.raise_for_status()
            content = response.text[:20000]  # Limit per filing to manage token count
            combined_content += f"\n\nFiling {filing['accession_number']} ({filing['form_type']} - {filing['entity_name']}):\n{content}"
        except Exception as e:
            combined_content += f"\n\nFiling {filing['accession_number']}: Error fetching content - {e}"

    # Limit total content to 50k chars to fit API constraints
    combined_content = combined_content[:50000]

    # Construct prompt for OpenAI
    prompt = f"""
Analyze the following content to identify which companies are most likely to hire '{client_name}', 
a company/companies specializing in '{client_specialization}'. Do not mention anything about the SEC or EDGAR in the response. 
Review all provided filings and return a list of up to 20 companies ordered from most likely to least likely to invest 
in our client with: 1. Their name 2. A brief reason why they might invest (e.g., aligned interests, investment 
history) 3. A short paragraph describing each company. Focus on patterns like investment focus, industry alignment, 
or mentions related to '{client_specialization}'. 4. Look for the company's revenue in the files and pick the ones 
that are ideally between 5-30 million US dollars. If no clear candidates emerge, state that explicitly.
Content (up to 50,000 characters across {len(filings)} filings):
{combined_content}
"""

    try:
        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=3000,
            temperature=0.8
        )
        answer = response.choices[0].message.content.strip()
        return answer
    except Exception as e:
        return f"Error analyzing filings with OpenAI: {e}"


def get_user_input():
    """Get query keywords, client name, and specialization from the user."""
    query = input("Enter query keywords (e.g., 'fintech investment'): ").strip()
    client_names = input("Enter your client's company name or names: ").strip()
    client_specialization = input("Enter your client's specialization (e.g., 'AI-driven analytics'): ").strip()
    return query, client_names, client_specialization


def main():
    if not SEC_API_KEY or not OPENAI_API_KEY or "YOUR_" in [SEC_API_KEY, OPENAI_API_KEY]:
        print("Please set both SEC_API_KEY and OPENAI_API_KEY in the script.")
        return

    # Get user inputs
    query, client_name, client_specialization = get_user_input()
    if not query or not client_name or not client_specialization:
        print("Query keywords, client name, and specialization are required.")
        return

    # Fetch filings
    filings = fetch_filings(query, LIMIT)
    if not filings:
        print("No filings to analyze—stopping.")
        return

    # Display filings
    print("\nFilings to be analyzed:")
    for idx, filing in enumerate(filings, 1):
        print(f"\nFiling {idx}:")
        print(f"  Accession Number: {filing['accession_number']}")
        print(f"  Form Type: {filing['form_type']}")
        print(f"  Entity Name: {filing['entity_name']}")
        print(f"  Filing Date: {filing['filing_date']}")
        print(f"  Document URL: {filing['document_url']}")

    # Analyze with OpenAI and provide results
    print(
        f"\nAnalyzing potential investors for '{client_name}' specializing in '{client_specialization}' across {len(filings)} filings...")
    result = analyze_investment_potential_with_openai(filings, client_name, client_specialization)
    print(f"\nResults:\n{result}")

    doc.add_heading('US COMPANIES', level=2)
    doc.add_paragraph(result)
    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
