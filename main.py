import time
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.options import Options
from selenium.webdriver.firefox.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.firefox import GeckoDriverManager
from docx.shared import RGBColor
import os
from itertools import product

DOC_PATH = 'test.docx'

# Load existing document or create a new one
if os.path.exists(DOC_PATH):
    print(f"Loading existing document: {DOC_PATH}")
    doc = Document(DOC_PATH)
else:
    print(f"Creating new document: {DOC_PATH}")
    doc = Document()
    doc.add_heading('POTENTIAL INVESTORS', level=1)

# Your predefined combinations
dropdown_options = {
    "fund-country": [""],
    "sector-1": ["Real Estate"],
    "sector-2": ["Proptech"],
    "continent to invest in": ["America"],
    "country to invest in": ["United States"],
    "size of investment": ["<1m"],
    "currency": ["USD"],
    "min/maj": [""],
    "fund-type": ["VC", "PE", "LBO", "OPPORTUNIST"],
    "fund-specialty": ["Early", "Growth"]
}

# Generate all combinations
combinations = list(product(*dropdown_options.values()))
field_names = list(dropdown_options.keys())
print(f"Total combinations to process: {len(combinations)}")  # Should be 8

all_companies = set()
existing_headings = {}
current_level_2 = None
current_level_3 = None

# Load existing company data from document
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading'):
        level = int(paragraph.style.name.split()[1])
        if level == 2:
            current_level_2 = paragraph.text.strip()
            if current_level_2 not in existing_headings:
                existing_headings[current_level_2] = {}
        elif level == 3 and current_level_2:
            current_level_3 = paragraph.text.split(' (')[0].strip()
            if current_level_3 not in existing_headings[current_level_2]:
                existing_headings[current_level_2][current_level_3] = set()
        elif level == 4 and current_level_2 and current_level_3:
            company_name = paragraph.text.strip()
            existing_headings[current_level_2][current_level_3].add(company_name)
            all_companies.add(company_name)


def find_select_dropdown_by_label(driver, wait, label_text):
    """Find a select dropdown by its label text, using class-based selectors."""
    try:
        # Find label by text (case-insensitive)
        label_xpath = f"//label[translate(normalize-space(text()),'abcdefghijklmnopqrstuvwxyz'," \
                      f"'ABCDEFGHIJKLMNOPQRSTUVWXYZ')='{label_text.upper()}']"

        # Allow some flexibility in finding the label
        labels = driver.find_elements(By.XPATH, label_xpath)
        if not labels:
            # Try partial match if exact match fails
            label_xpath = f"//label[contains(translate(normalize-space(text()),'abcdefghijklmnopqrstuvwxyz'," \
                          f"'ABCDEFGHIJKLMNOPQRSTUVWXYZ'), '{label_text.upper()}')]"
            labels = driver.find_elements(By.XPATH, label_xpath)

        if not labels:
            print(f"No label found for '{label_text}'")
            return None

        # Find the closest select container
        for label in labels:
            # Try to find the select container in various ways

            # 1. Look for the select container directly following the label
            container = None

            # Try by class name from the HTML example
            try:
                # Using the class name pattern you provided
                container = label.find_element(By.XPATH,
                                               "following-sibling::div[contains(@class, 'select_baseStyles__') or "
                                               "contains(@class, 'select_white-no-border__')]")
            except NoSuchElementException:
                pass

            # If that didn't work, try to find any react-select container
            if not container:
                try:
                    container = label.find_element(By.XPATH,
                                                   "following-sibling::div[contains(@class, 'css-') and contains("
                                                   "@class, 'control')]")
                except NoSuchElementException:
                    pass

            # 2. Try parent approach
            if not container:
                try:
                    parent = label.find_element(By.XPATH, "./..")
                    container = parent.find_element(By.XPATH,
                                                    ".//div[contains(@class, 'select_baseStyles__') or contains(@class, 'css-') and contains(@class, 'control')]")
                except NoSuchElementException:
                    pass

            # If we found a container, get the input
            if container:
                try:
                    input_field = container.find_element(By.TAG_NAME, "input")
                    return input_field
                except NoSuchElementException:
                    continue

        # Last resort: look for any element that matches the class pattern
        selects = driver.find_elements(By.CSS_SELECTOR, "div.select_baseStyles__")
        if selects:
            for select in selects:
                try:
                    nearby_text = select.find_element(By.XPATH,
                                                      "./preceding-sibling::label[1] | "
                                                      "./parent::*/preceding-sibling::label[1]").text
                    if label_text.lower() in nearby_text.lower():
                        input_field = select.find_element(By.TAG_NAME, "input")
                        return input_field
                except NoSuchElementException:
                    continue

            # If we couldn't find a matching label, just return the first one as a fallback
            try:
                input_field = selects[0].find_element(By.TAG_NAME, "input")
                print(f"Using fallback selector for '{label_text}'")
                return input_field
            except NoSuchElementException:
                pass

        print(f"Could not find input for label '{label_text}' using class selectors")
        return None
    except Exception as e:
        print(f"Error finding select for '{label_text}': {e}")
        return None


# Map field names to more human-readable labels
field_to_label_map = {
    "fund-country": "Funds Location",
    "sector-1": "Target sector(s)",
    "sector-2": "Target specialty(ies)",
    "continent to invest in": "Target continent*",
    "country to invest in": "Target country*",
    "size of investment": "Ticket size(s)*",
    "currency": "USD",
    "min/maj": "Min/Maj",
    "fund-type": "Fund type(s)",
    "fund-specialty": "Funds specialty(ies)"
}


def auto_select_dropdown(wait, driver, field_name, input_value):
    """Select a dropdown option using class-based selectors with fallback mechanisms"""
    retries = 3
    while retries > 0:
        print(
            f"\n=== Processing {field_name} (Value: '{input_value}', Retry: {4 - retries}/3) ===")

        if not input_value or input_value == "":
            print(f"Leaving {field_name} blank as value is '{input_value}'.\n")
            return True

        try:
            # Try to find by class-based selectors
            label_text = field_to_label_map.get(field_name, field_name)
            input_field = find_select_dropdown_by_label(driver, wait, label_text)

            if not input_field:
                print(f"Trying generic approach for {field_name}")
                select_containers = driver.find_elements(By.CSS_SELECTOR,
                                                         "div[class*='select_baseStyles__'], div[class*='css-']["
                                                         "class*='control']")

                for container in select_containers:
                    try:
                        # Check placeholder text
                        placeholder = container.find_element(By.CSS_SELECTOR, "div[class*='placeholder']")
                        if label_text.lower() in placeholder.text.lower():
                            input_field = container.find_element(By.TAG_NAME, "input")
                            break
                    except NoSuchElementException:
                        continue

            if not input_field:
                print(f"Could not find input field for {field_name}. Skipping...")
                retries -= 1
                if retries > 0:
                    print(f"Retrying ({retries} attempts left)...")
                    time.sleep(2)
                continue

            # Interact with the field
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_field)
            time.sleep(0.5)

            # Try clicking on the container first (sometimes more reliable than clicking directly on input)
            try:
                container = input_field.find_element(By.XPATH, "./ancestor::div[contains(@class, 'control')]")
                driver.execute_script("arguments[0].click();", container)
            except:
                input_field.click()

            time.sleep(0.5)

            for _ in range(2):
                input_field.clear()
                time.sleep(0.5)
                if not input_field.get_attribute('value'):
                    break

            input_field.send_keys(input_value)
            time.sleep(1)

            # Find and select the matching option
            options = wait.until(
                EC.presence_of_all_elements_located((By.CSS_SELECTOR,
                                                     "div[role='option'], div[class*='option'], div[class*='menu'] div")),
                message=f"Dropdown options not found for {field_name}"
            )

            found_option = None
            for option in options:
                option_text = option.text.strip()
                if input_value.lower() in option_text.lower():
                    found_option = option
                    break

            if found_option:
                driver.execute_script("arguments[0].scrollIntoView(true);", found_option)
                try:
                    found_option.click()
                except:
                    # If direct click fails, try JavaScript click
                    driver.execute_script("arguments[0].click();", found_option)

                time.sleep(1)
                print(f"Selected '{input_value}' for {field_name}.\n")
                return True
            else:
                print(f"No matching option found for {field_name} with value '{input_value}'.")
                retries -= 1

            if retries > 0:
                print(f"Retrying ({retries} attempts left)...")
                time.sleep(2)
            else:
                print(f"Max retries reached; skipping {field_name}.")
                return True

        except TimeoutException as e:
            print(f"Timeout error: {e}")
            retries -= 1
            if retries > 0:
                print(f"Retrying ({retries} attempts left)...")
                time.sleep(2)
            else:
                print(f"Max retries reached; skipping {field_name}.")
                return True
        except Exception as e:
            print(f"Unexpected error for {field_name}: {e}")
            retries -= 1
            if retries > 0:
                print(f"Retrying ({retries} attempts left)...")
                time.sleep(2)
            else:
                print(f"Max retries reached; skipping {field_name}.")
                return True


def insert_under_heading(doc, level_2, level_3, company_name, collected_data, num_results=None):
    """Insert company data under appropriate headings in the document"""
    found_level_2 = False
    found_level_3 = False
    insert_position = None

    if level_2 not in existing_headings:
        existing_headings[level_2] = {}
    if level_3 not in existing_headings[level_2]:
        existing_headings[level_2][level_3] = set()

    first_in_category = company_name not in existing_headings[level_2][level_3]
    company_exists = company_name in all_companies
    header_text = f"{level_3} ({num_results if num_results is not None else 0})"

    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 2':
            if para.text.strip() == level_2:
                found_level_2 = True
                insert_position = i + 1
            elif found_level_2:
                break
        elif found_level_2 and para.style.name == 'Heading 3' and para.text.split(' (')[0].strip() == level_3:
            found_level_3 = True
            para.text = header_text
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
            insert_position = i + 1

    if not found_level_2:
        doc.add_heading(level_2, level=2)
        new_heading = doc.add_heading(header_text, level=3)
        for run in new_heading.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
        new_company = doc.add_heading(company_name, level=4)
        if not company_exists:
            for item in reversed(collected_data):
                cleaned_item = item.rstrip('\n')
                if any(keyword in cleaned_item.lower() for keyword in ["linkedin url", "website url", "email"]):
                    item_with_space = cleaned_item + " "
                    paragraph = doc.add_paragraph(item_with_space, style='List Bullet')
                else:
                    paragraph = doc.add_paragraph(cleaned_item, style='List Bullet')
                new_company._element.addnext(paragraph._element)
        else:
            print(f"Company '{company_name}' already exists in document; appended name only.")
        existing_headings[level_2][level_3].add(company_name)
        all_companies.add(company_name)
    elif not found_level_3:
        new_heading = doc.add_heading(header_text, level=3)
        for run in new_heading.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
        doc.paragraphs[insert_position - 1]._element.addnext(new_heading._element)
        new_company = doc.add_heading(company_name, level=4)
        new_heading._element.addnext(new_company._element)
        if not company_exists:
            for item in reversed(collected_data):
                cleaned_item = item.rstrip('\n')
                if any(keyword in cleaned_item.lower() for keyword in ["linkedin url", "website url", "email"]):
                    item_with_space = cleaned_item + " "
                    paragraph = doc.add_paragraph(item_with_space, style='List Bullet')
                else:
                    paragraph = doc.add_paragraph(cleaned_item, style='List Bullet')
                new_company._element.addnext(paragraph._element)
        else:
            print(f"Company '{company_name}' already exists in document; appended name only.")
        existing_headings[level_2][level_3].add(company_name)
        all_companies.add(company_name)
    else:
        new_company = doc.add_heading(company_name, level=4)
        doc.paragraphs[insert_position - 1]._element.addnext(new_company._element)
        if not company_exists:
            for item in reversed(collected_data):
                cleaned_item = item.rstrip('\n')
                if any(keyword in cleaned_item.lower() for keyword in ["linkedin url", "website url", "email"]):
                    item_with_space = cleaned_item + " "
                    paragraph = doc.add_paragraph(item_with_space, style='List Bullet')
                else:
                    paragraph = doc.add_paragraph(cleaned_item, style='List Bullet')
                new_company._element.addnext(paragraph._element)
            existing_headings[level_2][level_3].add(company_name)
        else:
            print(f"Company '{company_name}' already exists in document; appended name only.")
        existing_headings[level_2][level_3].add(company_name)
        all_companies.add(company_name)


# Setup Selenium
options = Options()
options.headless = True  # Run in headless mode
driver_path = GeckoDriverManager().install()
if not os.access(driver_path, os.X_OK):
    os.chmod(driver_path, 0o755)  # Make the driver executable
service = Service(executable_path=driver_path, port=4444, log_path="geckodriver.log")
driver = webdriver.Firefox(service=service, options=options)
URL = "https://www.fundssniper.com/en/search-funds"

# Main loop over combinations
for i, combo in enumerate(combinations, 1):
    user_inputs = dict(zip(field_names, combo))
    print(f"\nProcessing combination {i}/{len(combinations)}: {user_inputs}")

    try:
        driver.get(URL)
        wait = WebDriverWait(driver, 15)
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, ".grow")))

        # Process each field in the form
        for field_name, input_value in user_inputs.items():
            success = auto_select_dropdown(wait, driver, field_name, input_value)
            if not success:
                print(f"Failed to process {field_name}; skipping this combination.")
                break

        # Submit the form
        time.sleep(2)
        submit_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button[data-variant='primary']")))
        submit_button.click()
        print("Search button clicked.")
        time.sleep(4)

        # Get number of results
        try:
            results_element = wait.until(EC.presence_of_element_located(
                (By.CSS_SELECTOR, "p.font-roboto.font-semibold.text-subtitle1.m-0.undefined")))
            results_text = results_element.text.strip()
            digits = ''.join(filter(str.isdigit, results_text))
            num_results = int(digits) if digits else 0
            print(f"Number of Results: {num_results}")
        except TimeoutException:
            num_results = 0
            print("Could not find number of results.")

        # Extract category information from user inputs
        sector = user_inputs.get("sector-1", "Unknown")
        fund_type = user_inputs.get("fund-type", "Unknown").upper()
        fund_specialty = user_inputs.get("fund-specialty", "Unknown").upper()
        category = f"{sector}-{fund_type}-{fund_specialty}"

        # Process View buttons
        original_window = driver.current_window_handle

        while True:
            # Find non-blurred view buttons
            view_buttons = driver.find_elements(
                By.XPATH,
                "//button[normalize-space(text())='View' and not(ancestor::div[contains(@class, 'absolute') and "
                "contains(@class, 'z-20') and contains(@class, 'flex') and contains(@class, 'justify-center')])]"
            )
            visible_buttons = [btn for btn in view_buttons if btn.is_displayed()]

            if not visible_buttons:
                print("No more non-blurred view buttons found.")
                break

            view_button = visible_buttons[0]
            print("Clicking a non-blurred view button...")
            driver.execute_script("arguments[0].scrollIntoView(true);", view_button)
            time.sleep(1)

            # Click the button and handle new window if opened
            old_windows = driver.window_handles.copy()
            driver.execute_script("arguments[0].click();", view_button)

            try:
                wait_long = WebDriverWait(driver, 10)
                wait_long.until(lambda d: len(d.window_handles) > len(old_windows))
                new_window = next(handle for handle in driver.window_handles if handle not in old_windows)
                driver.switch_to.window(new_window)
                print("New window detected.")
            except TimeoutException:
                print("No new window detected; detail page loaded in the same tab.")

            # Wait for detail page to load
            try:
                wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "h1")))
            except TimeoutException:
                print("Detail page did not load in time.")

            # Extract company information
            detail_html = driver.page_source
            detail_soup = BeautifulSoup(detail_html, 'html.parser')

            # Company name and basic info
            specific_h4_company_name = detail_soup.find('h4', class_="font-kodchasan font-light text-h8 undefined")
            specific_h4_about = detail_soup.find('h4', class_="font-kodchasan font-medium leading-[120%] text-h13 "
                                                              "xl:text-h4 mb-[1rem]")
            specific_p_about = detail_soup.find('p', class_="font-roboto font-regular text-subtitle3 m-0 undefined")

            # Main information section
            main_container = detail_soup.find("div", class_="2xl:max-w-[1264px] 2xl:mx-auto xl:mx-[50px] lg:mx-[44px] "
                                                            "md:mx-[40px] sm:mx-[16px] mx-[8px]")
            main_title = ""
            main_texts = ""
            if main_container:
                main_tags = main_container.select_one("p.font-roboto.font-semibold.text-subtitle1")
                main_info = main_container.select("p.font-roboto.font-semibold.text-subtitle2.m-0.undefined")
                main_title = main_tags.get_text(strip=True) if main_tags else ""
                main_texts = ", ".join(tag.get_text(strip=True) for tag in main_info)

            # Other sections
            containers = detail_soup.find_all("div", class_="flex flex-col items-center justify-between lg:flex-row")
            grouped_info = []
            for container in containers:
                header_tag = container.select_one("p.font-roboto.font-semibold.text-subtitle1")
                info_tags = container.select("p.font-roboto.font-semibold.text-subtitle2")
                if header_tag:
                    header_text = header_tag.get_text(strip=True)
                    info_texts = ", ".join(tag.get_text(strip=True) for tag in info_tags)
                    grouped_info.append(f"{header_text}: {info_texts}")

            # LinkedIn URL
            company_linkedin_url = ""
            try:
                company_linkedin_element = WebDriverWait(driver, 5).until(
                    EC.presence_of_element_located((
                        By.XPATH,
                        "//a[contains(., 'LinkedIn') and not(ancestor::div[contains(., 'team') or contains(., 'Team')])]"
                    ))
                )
                company_linkedin_url = company_linkedin_element.get_attribute("href")
                print(f"Found company LinkedIn URL: {company_linkedin_url}")
            except TimeoutException:
                print("No company LinkedIn link found outside team section.")

            # Team information
            collected_data = []
            has_team_section = detail_soup.find(
                lambda tag: tag.name in ['h2', 'h3', 'div'] and 'team' in tag.text.lower()
            )
            team_info = []
            if has_team_section:
                print("Team section detected; scraping team info.")
                team_containers = detail_soup.find_all('div', class_="flex items-center")
                if not team_containers:
                    team_containers = [has_team_section]

                for container in team_containers:
                    p_tags = container.find_all('p', class_="font-roboto font-regular text-subtitle3 m-0 mb-[1rem]")
                    name = p_tags[0].get_text(strip=True) if p_tags else ""
                    title = p_tags[1].get_text(strip=True) if len(p_tags) > 1 else ""
                    email_link = p_tags[3].find('a', href=lambda x: x and 'mailto:' in x) if len(p_tags) > 3 else None
                    email = email_link.get_text(strip=True) if email_link else ""
                    linkedin_link = container.find('a', href=lambda x: x and 'linkedin.com' in x)
                    linkedin_url = linkedin_link['href'] if linkedin_link and linkedin_link.get(
                        'href') != company_linkedin_url else ""

                    if not name:
                        continue
                    formatted_entry = f"{name}"
                    if title:
                        formatted_entry += f"\n  {title}"
                    if linkedin_url:
                        formatted_entry += f"\n  LinkedIn: {linkedin_url}"
                    if email:
                        formatted_entry += f"\n  Email: {email}"
                    print(f"Team member:\n{formatted_entry}\n")
                    team_info.append(formatted_entry)

                full_team = "\n\n".join(team_info) if team_info else "No team members found"
            else:
                print("No team section found; skipping team info scraping.")
                full_team = ""

            # Address information
            address_parts = [p.get_text(strip=True) for p in detail_soup.find_all(
                "p", class_="font-roboto font-regular text-subtitle3 m-0 text-center"
            ) if p.get_text(strip=True) and "ADDRESS" not in p.get_text()]
            full_address = ", ".join(address_parts)

            # Website URL
            website_href = ""
            try:
                website_element = WebDriverWait(driver, 5).until(
                    EC.presence_of_element_located((By.XPATH, "//a[contains(., 'WEBSITE')]"))
                )
                website_href = website_element.get_attribute("href")
            except TimeoutException:
                print("Website link not found; setting as empty.")

            # Contact information
            email_label = next((p for p in detail_soup.find_all("p") if "EMAIL" in p.get_text(strip=True).upper()),
                               None)
            call_label = next((p for p in detail_soup.find_all("p") if "CALL" in p.get_text(strip=True).upper()), None)
            email_address = email_label.find_next_sibling("a").find("p").get_text(
                strip=True) if email_label and email_label.find_next_sibling("a") else ""
            call_address = call_label.find_next_sibling("a").get_text(
                strip=True) if call_label and call_label.find_next_sibling("a") else ""

            # Compile all collected data
            if specific_h4_company_name:
                collected_data.append("Company: " + specific_h4_company_name.get_text(strip=True))
            if specific_h4_about:
                collected_data.append("About: " + specific_h4_about.get_text(separator=" ", strip=True))
            if specific_p_about:
                collected_data.append("Description: " + specific_p_about.get_text(strip=True))
            if main_title:
                collected_data.append("Main Info: " + f"{main_title}: {main_texts}")
            if grouped_info:
                collected_data.extend(grouped_info)
            if full_team.strip():
                collected_data.append("Team: \n" + full_team)
            if full_address:
                collected_data.append("Address: " + full_address)
            if company_linkedin_url:
                collected_data.append(f"Company LinkedIn URL: {company_linkedin_url}")
            if website_href:
                collected_data.append(f"Website URL: {website_href}")
            if email_address:
                collected_data.append(f"Email: {email_address}")
            if call_address:
                collected_data.append(f"Call: {call_address}")

            # Get fund information for categories
            country = user_inputs.get("fund-country", "Unknown").upper()
            sector1 = user_inputs.get("sector-1", "Unknown")
            sector2 = user_inputs.get("sector-2", "Unknown")
            fund_type = user_inputs.get("fund-type", "Unknown").upper()
            fund_specialty = user_inputs.get("fund-specialty", "Unknown").upper()

            fund_types = ["VC", "PE", "LBO", "OPPORTUNIST"]
            fund_specialties = ["INCUBATION", "SEED", "EARLY", "GROWTH"]
            fund_type = fund_type if fund_type in fund_types else "Unknown"
            fund_specialty = fund_specialty if fund_specialty in fund_specialties else "Unknown"

            us_category = f"US-FUNDS-{sector1}-{sector2}-{fund_type}-{fund_specialty}"
            all_category = f"{sector1}-{sector2}-{fund_type}-{fund_specialty}"

            company_name = specific_h4_company_name.get_text(
                strip=True) if specific_h4_company_name else "Unknown Company"

            # Insert data into appropriate document sections
            if country == "US":
                insert_under_heading(doc, "US FUNDS", us_category, company_name, collected_data, num_results)
            insert_under_heading(doc, "ALL FUNDS", all_category, company_name, collected_data, num_results)

            print(f"Processed {company_name} for categories: {us_category if country == 'US' else ''} {all_category}")

            # Return to results page
            if len(driver.window_handles) > len(old_windows):
                driver.close()
                driver.switch_to.window(old_windows[0])
            else:
                driver.back()

            # Remove the processed button to avoid clicking it again
            driver.execute_script("arguments[0].remove();", view_button)
            time.sleep(0.5)

        # Save document after processing each combination
        doc.save(DOC_PATH)
        print(f"Updated document saved: {DOC_PATH}")

    except Exception as e:
        print(f"Error processing combination {user_inputs}: {e}")
        continue

# Clean up and save final document
driver.quit()
doc.save(DOC_PATH)
print(f"Program completed. Final document saved: {DOC_PATH}")