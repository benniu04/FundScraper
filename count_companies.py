from docx import Document
import os
from collections import Counter

DOC_PATH = 'Aerospace.docx'

def count_and_rank_companies(doc_path):
    if not os.path.exists(doc_path):
        print(f"Error: The file {doc_path} does not exist.")
        return None, 0

    doc = Document(doc_path)

    company_counts = Counter()
    total_entries = 0

    for para in doc.paragraphs:
        if para.style.name == 'Heading 4':
            company_name = para.text.strip()
            company_counts[company_name] += 1
            total_entries += 1

    repeated_companies = []
    for name, count in company_counts.most_common():
        if count > 1:
            repeated_companies.append((name, count))

    print(f"Total company entries: {total_entries}")
    print(f"Unique companies: {len(company_counts)}")

    return repeated_companies, total_entries


def update_document_with_ranking(doc_path, repeated_companies, total_entries):
    doc = Document(doc_path)
    new_doc = Document()

    new_doc.add_heading("POTENTIAL INVESTORS", level=1)
    new_doc.add_heading("ALL FUNDS", level=2)
    new_doc.add_paragraph(f"Number of Results: {total_entries}")

    if repeated_companies:
        new_doc.add_paragraph("A ranking of how many times a company is being repeated: ")
        for i, (company, count) in enumerate(repeated_companies, 1):
            new_doc.add_paragraph(f"{company} - {count} occurrences", style='List Number')
    else:
        new_doc.add_paragraph("No companies appear more than once.")

    new_doc.add_paragraph("")

    for element in doc.element.body:
        new_doc.element.body.append(element)

    new_doc.save(doc_path)
    print(f"\nDocument updated with ranking saved: {doc_path}")

if __name__ == "__main__":
    repeated_companies, total_entries = count_and_rank_companies(DOC_PATH)

    if total_entries == 0:
        print("No companies found or unable to process the document.")
    elif repeated_companies is None:
        print("No companies found or unable to process the document.")
    else:
        update_document_with_ranking(DOC_PATH, repeated_companies, total_entries)

